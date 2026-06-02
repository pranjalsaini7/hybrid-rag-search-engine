"""
Document Loaders — PDF, DOCX, TXT

Each loader returns a list[Document] with metadata:
  {source: filename, page: int, file_type: str}

Uses PyMuPDF (fitz) for PDFs because it preserves layout better
than PyPDF2 and handles scanned-text PDFs more gracefully.
"""

from __future__ import annotations

from pathlib import Path
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxFile
from langchain_core.documents import Document


def load_pdf(file_path: str | Path) -> List[Document]:
    """
    Extract text from a PDF page-by-page.

    Each page becomes one Document so that page numbers are preserved
    in metadata for downstream source attribution.
    """
    file_path = Path(file_path)
    documents: list[Document] = []

    with fitz.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if not text:
                continue
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path.name,
                        "page": page_num,
                        "file_type": "pdf",
                        "file_path": str(file_path),
                    },
                )
            )

    return documents


def load_docx(file_path: str | Path) -> List[Document]:
    """
    Extract text from a DOCX file.

    DOCX files don't have page numbers in the same sense as PDFs,
    so we track paragraph index as a proxy for location.
    """
    file_path = Path(file_path)
    doc = DocxFile(str(file_path))

    full_text = "\n\n".join(
        para.text.strip() for para in doc.paragraphs if para.text.strip()
    )

    if not full_text:
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": file_path.name,
                "page": 1,
                "file_type": "docx",
                "file_path": str(file_path),
            },
        )
    ]


def load_txt(file_path: str | Path) -> List[Document]:
    """Load a plain text file as a single Document."""
    file_path = Path(file_path)
    text = file_path.read_text(encoding="utf-8", errors="replace").strip()

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": file_path.name,
                "page": 1,
                "file_type": "txt",
                "file_path": str(file_path),
            },
        )
    ]


# ── Dispatcher ──────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_txt,
}


def load_document(file_path: str | Path) -> List[Document]:
    """
    Auto-detect file type and load accordingly.

    Raises:
        ValueError: If the file extension is not supported.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    return loader(file_path)
